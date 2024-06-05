from docx import Document
from docx.shared import RGBColor

def change_color_one_letter_at_a_time(doc, baudot_sequence):
    index = 0  # Индекс текущей буквы в последовательности
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Читаем текст из run
            text = run.text
            # Создаем новый run для каждого символа и меняем цвет
            for char in text:
                if index < len(baudot_sequence) and char != ' ':
                    new_run = paragraph.add_run(char)
                    if baudot_sequence[index] == '1':
                        # Меняем цвет на 1.1.1
                        new_run.font.color.rgb = RGBColor(1, 1, 1)
                    elif baudot_sequence[index] == '0':
                        # Меняем цвет на 0.0.0
                        new_run.font.color.rgb = RGBColor(0, 0, 0)
                    index += 1  # Переходим к следующей букве в последовательности
                    # Если достигнут конец последовательности, выходим из цикла
                    if index >= len(baudot_sequence):
                        return

# Пример использования
doc = Document('1.docx')
baudot_sequence = '0000110010110001001100000101111000100010101100011110010110001111000101000001010001110100100011'  # Пример заданной последовательности
change_color_one_letter_at_a_time(doc, baudot_sequence)
doc.save('modified_document.docx')
